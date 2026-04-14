"""Generate summary .docx for team sharing."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading("One-Click Report Analysis Agent: Build Summary", level=1)

# Summary
doc.add_heading("Summary", level=2)
p = doc.add_paragraph()
p.add_run(
    "I built an AWS Bedrock Agent that automates One-Click Report root-cause analysis, "
    "replacing the manual process of cross-referencing OneTrack exports, Vlocity Error Logs, "
    "and PS Exception Logs in Salesforce. "
)
p.add_run(
    "A REST API (POST /analyze) accepts report data, the Bedrock Agent orchestrates four "
    "Lambda action groups to investigate, and returns a structured analysis with root cause, "
    "severity, and recommended next steps. "
)
p.add_run(
    "Salesforce is mocked with DynamoDB for now, designed to swap to live SF API when ready."
)

# Lambda details
doc.add_heading("What Each Lambda Does", level=2)

lambdas = [
    ("ParseReport", "Extracts key fields from the report and pulls Salesforce record IDs out of error messages for direct log lookup."),
    ("QueryVlocityLogs", "Retrieves Vlocity Error Logs by ID or searches by agent LAN ID + timeframe, then parses HTTP request/response payloads to surface error details."),
    ("QueryExceptionLogs", "Looks up PS Exception Logs by ID to get exception type, severity, location, and error message."),
    ("ClassifyIssue", 'Categorizes vague agent descriptions (e.g. "SCREEN FROZE", "SPINNING") into LATENCY, UI_ERROR, AUTH_ERROR, DATA_ERROR, or UNKNOWN, and flags when recording review is needed.'),
    ("Entrypoint", "Receives the API request, invokes the Bedrock Agent, and returns the final JSON analysis + human-readable summary."),
]

for name, desc in lambdas:
    p = doc.add_paragraph(style="List Bullet")
    bold_run = p.add_run(name)
    bold_run.bold = True
    p.add_run(f": {desc}")

# Also created
doc.add_heading("Also Created", level=2)

extras = [
    "OpenAPI schemas for each action group (4 files)",
    "Agent instructions prompt defining the full 5-step analysis workflow",
    "Mock data derived from the original One-Click Analysis document screenshots",
    "Full Terraform infrastructure (API Gateway, Bedrock Agent, DynamoDB, S3, IAM)",
    "35 passing unit tests covering all Lambda handlers",
]

for item in extras:
    doc.add_paragraph(item, style="List Bullet")

# Architecture
doc.add_heading("Architecture", level=2)
p = doc.add_paragraph()
p.style.font.name = "Consolas"
arch = (
    "POST /analyze\n"
    "    │\n"
    "    ▼\n"
    "API Gateway (REST)\n"
    "    │\n"
    "    ▼\n"
    "Entrypoint Lambda --> Bedrock Agent (Claude 3.5 Sonnet)\n"
    "                           │\n"
    "                   ┌───────┼────────────┬────────────────┐\n"
    "                   ▼       ▼            ▼                ▼\n"
    "             ParseReport  QueryVlocity  QueryException  ClassifyIssue\n"
    "             (Lambda)     Logs(Lambda)  Logs(Lambda)    (Lambda)\n"
    "                   │       │            │\n"
    "                   ▼       ▼            ▼\n"
    "                 S3      DynamoDB     DynamoDB\n"
    "              (reports)  (mock SF)   (mock SF)"
)
run = p.add_run(arch)
run.font.name = "Consolas"
run.font.size = Pt(9)

# Next steps
doc.add_heading("Next Steps", level=2)
next_steps = [
    "Deploy via Terraform (terraform init > plan > apply)",
    "Test with sample One-Click Report data against the /analyze endpoint",
    "Swap mock DynamoDB layer for live Salesforce API (Connected App + Secrets Manager)",
    "Add Bedrock Knowledge Base with OmniScript docs for smarter analysis",
    "Add batch processing (S3 trigger → Step Functions) for full Excel exports",
]
for item in next_steps:
    doc.add_paragraph(item, style="List Bullet")

doc.save(r"c:\Users\dzgw\Desktop\Oneclick agent\OneClick Agent Build Summary.docx")
print("Done")
